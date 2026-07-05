"""
JobHunt LatAm — remote-job search + resume matching web app.

Features
  1) Aggregates remote jobs from free public APIs (Remotive, RemoteOK,
     Arbeitnow, Jobicy, Himalayas, The Muse, We Work Remotely, Working Nomads,
     Jobspresso) and flags which ones a LatAm-based applicant can realistically
     take given their country (worldwide / LatAm / Spain only).
  2) Reads an uploaded resume (.docx, .pdf or .txt) and ranks jobs by fit,
     using ONLY the work-experience section (ignores courses/education).
  3) Suggests resume changes tailored to a specific role.
  4) "Apply changes" builds a new tailored .docx you can download.
  5) Optional salary filter + per-country eligibility filter.

Matching + suggestions work fully offline (local scoring). If you paste an
OpenAI or Anthropic API key in Settings, the app upgrades to LLM-quality
matching and rewrite suggestions.
"""

import io
import os
import re
import json
import time
import html
import sys
import subprocess
import importlib
import threading
import xml.etree.ElementTree as ET
from datetime import datetime


def _ensure(module, pip_name=None):
    """Import a module, auto-installing it into THIS interpreter if missing."""
    try:
        return importlib.import_module(module)
    except ImportError:
        pip_name = pip_name or module
        print("[setup] Installing missing dependency: %s ..." % pip_name, flush=True)
        try:
            subprocess.check_call(
                [sys.executable, "-m", "pip", "install", "--quiet", pip_name])
        except Exception as e:
            print("[setup] Could not auto-install %s (%s). "
                  "Please run: %s -m pip install %s"
                  % (pip_name, e, sys.executable, pip_name), flush=True)
            return None
        try:
            return importlib.import_module(module)
        except ImportError:
            return None


_ensure("flask", "Flask")
_ensure("requests", "requests")
docx = _ensure("docx", "python-docx")
pdfplumber = _ensure("pdfplumber", "pdfplumber")

import requests
from flask import (
    Flask, request, jsonify, send_file, render_template, Response
)

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16 MB uploads

# ---------------------------------------------------------------------------
# In-memory state (single-user local app)
# ---------------------------------------------------------------------------
STATE = {
    "resume_text": "",          # full resume text (used for export)
    "resume_experience": "",    # ONLY the work-experience section (used for matching)
    "resume_docx": None,        # original .docx bytes (to preserve formatting on apply)
    "resume_filename": "",
    "resume_skills": [],
    "last_jobs": [],
    "settings": {
        "llm_provider": "none",
        "api_key": "",
        "openai_model": "gpt-4o-mini",
        "anthropic_model": "claude-3-5-sonnet-latest",
    },
}
_LOCK = threading.Lock()

HTTP_HEADERS = {
    "User-Agent": "Mozilla/5.0 (JobHunt-LatAm/2.0; +local app)",
    "Accept": "application/json, text/xml, */*",
}

# ---------------------------------------------------------------------------
# Location / eligibility model
# ---------------------------------------------------------------------------
# LatAm countries the user could be based in (the dropdown is built from these).
LATAM_COUNTRIES = [
    "argentina", "bolivia", "brazil", "chile", "colombia", "costa rica",
    "cuba", "dominican republic", "ecuador", "el salvador", "guatemala",
    "honduras", "mexico", "nicaragua", "panama", "paraguay", "peru",
    "puerto rico", "uruguay", "venezuela",
]
# Alternate spellings mapped to the canonical name above.
COUNTRY_ALIASES = {
    "brasil": "brazil", "méxico": "mexico", "perú": "peru",
}

# "Whole region" signals → eligible for any LatAm applicant.
LATAM_WIDE = [
    "latam", "lat-am", "latin america", "latinoamerica", "latinoamérica",
    "south america", "central america", "the americas", "americas",
    "north and south america", "pan-american", "remote - americas",
    "remote (americas)", "americas time", "americas timezone",
    "gmt-3", "gmt-4", "gmt-5", "gmt-6",
    "utc-3", "utc-4", "utc-5", "utc-6", "utc-03", "utc-04", "utc-05", "utc-06",
]
# Strong "anywhere in the world" signals → eligible for everyone.
STRONG_GLOBAL = [
    "worldwide", "work from anywhere", "anywhere in the world", "global remote",
    "any country", "no location requirement", "location independent",
    "remote - worldwide", "remote (worldwide)", "work from any country",
    "work from anywhere in the world",
]
# Specific non-LatAm places that imply you must already have local work
# authorization there → NOT eligible (you have no visa).  Spain is handled
# separately because it sometimes sponsors.
NONLATAM_LOCATIONS = [
    "united states", "u.s.a", "u.s.", "usa", "us-based", "us based",
    "us only", "americas-us", "canada", "canadian",
    "united kingdom", "u.k.", "uk only", "england", "scotland", "ireland",
    "portugal", "germany", "deutschland", "france", "italy", "netherlands",
    "belgium", "switzerland", "austria", "poland", "romania", "czech",
    "sweden", "norway", "denmark", "finland", "greece", "hungary",
    "europe", "european union", "eu only", "emea", "uk/eu", "eu/uk",
    "apac", "asia", "india", "pakistan", "bangladesh", "philippines",
    "vietnam", "indonesia", "australia", "new zealand", "singapore",
    "malaysia", "japan", "china", "hong kong", "south korea",
    "africa", "nigeria", "kenya", "egypt", "south africa",
    "middle east", "gcc", "uae", "dubai", "saudi", "israel", "turkey",
]
# Definite hard "no" phrases.
NEGATIVE_GEO = [
    "us only", "u.s. only", "usa only", "united states only", "us-only",
    "must be located in the us", "must reside in the united states",
    "us residents only", "based in the us", "must be us based",
    "authorized to work in the us", "eu only", "europe only", "uk only",
    "emea only", "must be based in europe", "canada only", "india only",
    "apac only", "australia only", "must be authorized to work in",
]
SPAIN_TERMS = ["spain", "españa", "espana", "madrid", "barcelona", "valencia"]


def _norm_country(c):
    c = (c or "all").strip().lower()
    return COUNTRY_ALIASES.get(c, c)


def latam_eligibility(text, region, country="all"):
    """Classify a job for a LatAm-based applicant.

    Returns (label, score 0-100, reason). Higher = more eligible.
    `country` is 'all' or a specific LatAm country the user is based in.
    Rule of thumb: only worldwide/anywhere, LatAm-wide, the user's own LatAm
    country, or Spain (possible sponsorship) count as eligible. Jobs that are
    "remote *from* USA / Portugal / Europe / etc." are marked NOT eligible
    because they require local work authorization the user doesn't have.
    """
    sel = _norm_country(country)
    blob = f"{region} {text}".lower()

    neg = next((p for p in NEGATIVE_GEO if p in blob), None)
    latam_wide = next((p for p in LATAM_WIDE if p in blob), None)
    found = [c for c in LATAM_COUNTRIES if c in blob]
    spain = any(s in blob for s in SPAIN_TERMS)
    strong_global = any(g in blob for g in STRONG_GLOBAL)
    anywhere_generic = ("anywhere" in blob and "anywhere in" not in blob)
    worldwide = strong_global or anywhere_generic
    other = next((c for c in NONLATAM_LOCATIONS if c in blob), None)

    # 1) Region-wide LatAm / Americas → good for any LatAm applicant.
    if latam_wide:
        return ("Eligible (LatAm)", 100, f"region: {latam_wide}")

    # 2) Worldwide / anywhere → good for everyone.
    if worldwide and not neg:
        return ("Eligible (worldwide)", 90, "open worldwide / anywhere")

    # 3) Specific LatAm country mentioned.
    if found:
        if sel == "all" or sel in found:
            return ("Eligible (LatAm)", 100, f"mentions {found[0].title()}")
        # Restricted to a different LatAm country than the user's.
        return (f"Open to {found[0].title()} only", 25, "different LatAm country")

    # 4) Spain — sometimes sponsors, so allow it (only Spain among EU).
    if spain and not (other and other not in ("spain",)):
        return ("Eligible via Spain (may sponsor)", 70, "Spain-based role")

    # 5) Any other specific country/region → needs local work auth → no.
    if other or neg:
        return ("Not eligible (needs local work auth)", 8,
                f"requires {other or 'specific location'}")

    # 6) Generic remote with no stated geo → uncertain, show it.
    if "remote" in blob:
        return ("Possibly eligible", 50, "remote, no explicit geo")
    return ("Unknown", 35, "no location signal")


# ---------------------------------------------------------------------------
# Salary parsing
# ---------------------------------------------------------------------------
def _to_int(v):
    try:
        if v is None or v == "":
            return None
        return int(float(str(v).replace(",", "").replace("$", "").strip()))
    except (TypeError, ValueError):
        return None


def parse_salary(*texts):
    raw = " ".join(t for t in texts if t)
    if not raw:
        return None, None, ""
    candidates = re.findall(
        r"(?:USD|US\$|\$|EUR|€|£)?\s?(\d{1,3}(?:[,.\s]\d{3})+|\d{2,3})\s?[kK]?",
        raw,
    )
    nums = []
    for c in candidates:
        n = c.replace(",", "").replace(".", "").replace(" ", "")
        try:
            val = int(n)
        except ValueError:
            continue
        if val < 1000:
            val *= 1000
        if 8000 <= val <= 1_000_000:
            nums.append(val)
    if not nums:
        return None, None, ""
    return min(nums), max(nums), raw.strip()[:120]


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def _strip_html(s):
    if not s:
        return ""
    s = re.sub(r"<[^>]+>", " ", s)
    s = html.unescape(s)
    return re.sub(r"\s+", " ", s).strip()


def _kwlist(keywords):
    return [k.strip().lower() for k in re.split(r"[\s,]+", keywords or "") if k.strip()]


def _matches(keywords, *fields):
    kw = _kwlist(keywords)
    if not kw:
        return True
    hay = " ".join(str(f) for f in fields if f).lower()
    return any(k in hay for k in kw)


def _blank_job():
    return {
        "source": "", "title": "", "company": "", "region": "", "category": "",
        "tags": [], "url": "", "description": "", "posted": "",
        "salary_min": None, "salary_max": None, "salary_raw": "",
    }


# ---------------------------------------------------------------------------
# Job source adapters  (all free, no API key)
# ---------------------------------------------------------------------------
def fetch_remotive(keywords):
    out = []
    try:
        params = {"search": keywords, "limit": 60} if keywords else {"limit": 60}
        r = requests.get("https://remotive.com/api/remote-jobs", params=params,
                         headers=HTTP_HEADERS, timeout=20)
        r.raise_for_status()
        for j in r.json().get("jobs", []):
            desc = _strip_html(j.get("description", ""))
            smin, smax, sraw = parse_salary(j.get("salary", ""))
            jb = _blank_job()
            jb.update(source="Remotive", title=j.get("title", ""),
                      company=j.get("company_name", ""),
                      region=j.get("candidate_required_location", ""),
                      category=j.get("category", ""), tags=j.get("tags", []) or [],
                      url=j.get("url", ""), description=desc,
                      posted=j.get("publication_date", ""),
                      salary_min=smin, salary_max=smax,
                      salary_raw=sraw or j.get("salary", ""))
            out.append(jb)
    except Exception as e:
        app.logger.warning("Remotive failed: %s", e)
    return out


def fetch_remoteok(keywords):
    out = []
    try:
        r = requests.get("https://remoteok.com/api", headers=HTTP_HEADERS, timeout=20)
        r.raise_for_status()
        for j in r.json():
            if not isinstance(j, dict) or not j.get("position"):
                continue
            desc = _strip_html(j.get("description", ""))
            tags = j.get("tags", []) or []
            if not _matches(keywords, j.get("position", ""), desc, " ".join(tags)):
                continue
            smin = _to_int(j.get("salary_min"))
            smax = _to_int(j.get("salary_max"))
            sraw = f"${smin or '?'} - ${smax or '?'}" if (smin or smax) else ""
            jb = _blank_job()
            jb.update(source="RemoteOK", title=j.get("position", ""),
                      company=j.get("company", ""),
                      region=j.get("location", "") or "Remote",
                      category=", ".join(tags[:3]), tags=tags,
                      url=j.get("url", "") or j.get("apply_url", ""),
                      description=desc, posted=j.get("date", ""),
                      salary_min=smin, salary_max=smax, salary_raw=sraw)
            out.append(jb)
    except Exception as e:
        app.logger.warning("RemoteOK failed: %s", e)
    return out


def fetch_arbeitnow(keywords):
    out = []
    try:
        r = requests.get("https://www.arbeitnow.com/api/job-board-api",
                         headers=HTTP_HEADERS, timeout=20)
        r.raise_for_status()
        for j in r.json().get("data", []):
            desc = _strip_html(j.get("description", ""))
            tags = j.get("tags", []) or []
            if not _matches(keywords, j.get("title", ""), desc, " ".join(tags)):
                continue
            smin, smax, sraw = parse_salary(desc[:400])
            jb = _blank_job()
            jb.update(source="Arbeitnow", title=j.get("title", ""),
                      company=j.get("company_name", ""),
                      region=j.get("location", "") or ("Remote" if j.get("remote") else ""),
                      category=", ".join(tags[:3]), tags=tags,
                      url=j.get("url", ""), description=desc,
                      posted=str(j.get("created_at", "")),
                      salary_min=smin, salary_max=smax, salary_raw=sraw)
            out.append(jb)
    except Exception as e:
        app.logger.warning("Arbeitnow failed: %s", e)
    return out


def fetch_jobicy(keywords):
    out = []
    try:
        r = requests.get("https://jobicy.com/api/v2/remote-jobs",
                         params={"count": 50}, headers=HTTP_HEADERS, timeout=20)
        r.raise_for_status()
        for j in r.json().get("jobs", []):
            desc = _strip_html(j.get("jobDescription", "") or j.get("jobExcerpt", ""))
            tags = (j.get("jobIndustry", []) or []) + (j.get("jobType", []) or [])
            title = j.get("jobTitle", "")
            if not _matches(keywords, title, desc, " ".join(tags)):
                continue
            smin = _to_int(j.get("annualSalaryMin"))
            smax = _to_int(j.get("annualSalaryMax"))
            cur = j.get("salaryCurrency", "") or "$"
            sraw = f"{cur} {smin or '?'}-{smax or '?'}" if (smin or smax) else ""
            jb = _blank_job()
            jb.update(source="Jobicy", title=title,
                      company=j.get("companyName", ""),
                      region=j.get("jobGeo", "") or "Remote",
                      category=", ".join((j.get("jobIndustry", []) or [])[:2]),
                      tags=tags, url=j.get("url", ""), description=desc,
                      posted=str(j.get("pubDate", "")),
                      salary_min=smin, salary_max=smax, salary_raw=sraw)
            out.append(jb)
    except Exception as e:
        app.logger.warning("Jobicy failed: %s", e)
    return out


def fetch_himalayas(keywords):
    out = []
    try:
        r = requests.get("https://himalayas.app/jobs/api",
                         params={"limit": 50}, headers=HTTP_HEADERS, timeout=20)
        r.raise_for_status()
        for j in r.json().get("jobs", []):
            desc = _strip_html(j.get("description", "") or j.get("excerpt", ""))
            title = j.get("title", "")
            cats = j.get("categories", []) or []
            if not _matches(keywords, title, desc, " ".join(cats)):
                continue
            locs = j.get("locationRestrictions", []) or []
            tzs = j.get("timezones", []) or []
            region = ", ".join(locs) if locs else "Remote"
            if tzs:
                region += " (" + ", ".join(str(t) for t in tzs[:3]) + ")"
            smin = _to_int(j.get("minSalary"))
            smax = _to_int(j.get("maxSalary"))
            sraw = f"${smin or '?'}-${smax or '?'}" if (smin or smax) else ""
            jb = _blank_job()
            jb.update(source="Himalayas", title=title,
                      company=j.get("companyName", "") or j.get("company", ""),
                      region=region, category=", ".join(cats[:2]),
                      tags=cats, url=j.get("applicationLink", "") or j.get("url", ""),
                      description=desc, posted=str(j.get("pubDate", "")),
                      salary_min=smin, salary_max=smax, salary_raw=sraw)
            out.append(jb)
    except Exception as e:
        app.logger.warning("Himalayas failed: %s", e)
    return out


def fetch_themuse(keywords):
    out = []
    try:
        for page in (0, 1):
            r = requests.get("https://www.themuse.com/api/public/jobs",
                             params={"page": page, "location": "Flexible / Remote"},
                             headers=HTTP_HEADERS, timeout=20)
            r.raise_for_status()
            for j in r.json().get("results", []):
                title = j.get("name", "")
                desc = _strip_html(j.get("contents", ""))
                cats = [c.get("name", "") for c in (j.get("categories", []) or [])]
                if not _matches(keywords, title, desc, " ".join(cats)):
                    continue
                locs = [l.get("name", "") for l in (j.get("locations", []) or [])]
                jb = _blank_job()
                jb.update(source="The Muse", title=title,
                          company=(j.get("company") or {}).get("name", ""),
                          region=", ".join(locs) or "Remote",
                          category=", ".join(cats[:2]), tags=cats,
                          url=(j.get("refs") or {}).get("landing_page", ""),
                          description=desc, posted=str(j.get("publication_date", "")))
                out.append(jb)
    except Exception as e:
        app.logger.warning("The Muse failed: %s", e)
    return out


def _parse_rss(content):
    """Yield (title, link, description, region, pubdate) from an RSS feed."""
    root = ET.fromstring(content)
    for item in root.iter("item"):
        title = (item.findtext("title") or "").strip()
        link = (item.findtext("link") or "").strip()
        desc = _strip_html(item.findtext("description") or "")
        region = ""
        for ch in item:
            tag = ch.tag.split("}")[-1].lower()
            if tag in ("region", "location") and ch.text:
                region = ch.text.strip()
        yield title, link, desc, region, (item.findtext("pubDate") or "")


def fetch_weworkremotely(keywords):
    out = []
    try:
        r = requests.get("https://weworkremotely.com/remote-jobs.rss",
                         headers=HTTP_HEADERS, timeout=20)
        r.raise_for_status()
        for title, link, desc, region, pub in _parse_rss(r.content):
            company, sep, pos = title.partition(":")
            if sep:
                company, position = company.strip(), pos.strip()
            else:
                company, position = "", title
            if not _matches(keywords, title, desc):
                continue
            smin, smax, sraw = parse_salary(desc[:400])
            jb = _blank_job()
            jb.update(source="We Work Remotely", title=position or title,
                      company=company, region=region or "Remote", url=link,
                      description=desc, posted=pub,
                      salary_min=smin, salary_max=smax, salary_raw=sraw)
            out.append(jb)
    except Exception as e:
        app.logger.warning("We Work Remotely failed: %s", e)
    return out


def fetch_workingnomads(keywords):
    out = []
    try:
        r = requests.get("https://www.workingnomads.com/api/exposed_jobs/",
                         headers=HTTP_HEADERS, timeout=20)
        r.raise_for_status()
        data = r.json()
        rows = data if isinstance(data, list) else data.get("jobs", [])
        for j in rows:
            if not isinstance(j, dict):
                continue
            title = j.get("title", "")
            desc = _strip_html(j.get("description", ""))
            tags = j.get("tags", "")
            tags = tags.split(",") if isinstance(tags, str) else (tags or [])
            cat = j.get("category_name", "") or ""
            if not _matches(keywords, title, desc, cat, " ".join(tags)):
                continue
            smin, smax, sraw = parse_salary(desc[:400])
            jb = _blank_job()
            jb.update(source="Working Nomads", title=title,
                      company=j.get("company_name", "") or "",
                      region=j.get("location", "") or "Remote",
                      category=cat, tags=[t.strip() for t in tags if t],
                      url=j.get("url", ""), description=desc,
                      posted=str(j.get("pub_date", "")),
                      salary_min=smin, salary_max=smax, salary_raw=sraw)
            out.append(jb)
    except Exception as e:
        app.logger.warning("Working Nomads failed: %s", e)
    return out


def fetch_jobspresso(keywords):
    out = []
    try:
        r = requests.get("https://jobspresso.co/?feed=job_feed",
                         headers=HTTP_HEADERS, timeout=20)
        r.raise_for_status()
        for title, link, desc, region, pub in _parse_rss(r.content):
            # title is often "Position at Company" or "Company: Position"
            company, position = "", title
            if " at " in title:
                position, _, company = title.partition(" at ")
            elif ":" in title:
                company, _, position = title.partition(":")
            if not _matches(keywords, title, desc):
                continue
            smin, smax, sraw = parse_salary(desc[:400])
            jb = _blank_job()
            jb.update(source="Jobspresso", title=position.strip() or title,
                      company=company.strip(), region=region or "Remote",
                      url=link, description=desc, posted=pub,
                      salary_min=smin, salary_max=smax, salary_raw=sraw)
            out.append(jb)
    except Exception as e:
        app.logger.warning("Jobspresso failed: %s", e)
    return out


def fetch_getonbrd(keywords):
    """Get on Board (getonbrd.com) — LatAm-focused tech jobs, public JSON:API."""
    out = []
    try:
        params = {"per_page": 50}
        if keywords:
            params["query"] = keywords
        r = requests.get("https://www.getonbrd.com/api/v0/search/jobs",
                         params=params, headers=HTTP_HEADERS, timeout=20)
        r.raise_for_status()
        for item in r.json().get("data", []):
            a = item.get("attributes", {}) or {}
            title = a.get("title", "")
            if not title:
                continue
            desc = _strip_html(a.get("description", "") or a.get("functions", "")
                               or a.get("desirable", ""))
            cat = a.get("category_name", "") or ""
            if not _matches(keywords, title, desc, cat):
                continue
            remote = bool(a.get("remote")) or (a.get("remote_modality") in
                                               ("fully_remote", "partially_remote", "hybrid"))
            # Get on Board is a LatAm board: remote roles are open to LatAm.
            region = "Latin America (remote)" if remote else (a.get("country", "") or "")
            smin = _to_int(a.get("min_salary"))
            smax = _to_int(a.get("max_salary"))
            sraw = f"${smin or '?'}-${smax or '?'}" if (smin or smax) else ""
            links = item.get("links", {}) or {}
            url = links.get("public_url", "") or a.get("public_url", "")
            jb = _blank_job()
            jb.update(source="Get on Board", title=title,
                      company=a.get("company_name", "") or "",
                      region=region, category=cat, tags=[cat] if cat else [],
                      url=url, description=desc, posted=str(a.get("published_at", "")),
                      salary_min=smin, salary_max=smax, salary_raw=sraw)
            out.append(jb)
    except Exception as e:
        app.logger.warning("Get on Board failed: %s", e)
    return out


_HIRINGCAFE_URL = "https://hiring.cafe/api/search-jobs"


def _hiringcafe_headers():
    return {
        "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                       "AppleWebKit/537.36 (KHTML, like Gecko) "
                       "Chrome/130.0.0.0 Safari/537.36"),
        "Accept": "application/json, text/plain, */*",
        "Accept-Language": "en-US,en;q=0.9",
        "Content-Type": "application/json",
        "Origin": "https://hiring.cafe",
        "Referer": "https://hiring.cafe/",
    }


def _hiringcafe_full_state(keywords):
    return {
        "locations": [{
            "formatted_address": "United States",
            "types": ["country"],
            "geometry": {"location": {"lat": "39.8283", "lon": "-98.5795"}},
            "id": "user_country",
            "address_components": [{"long_name": "United States",
                                    "short_name": "US", "types": ["country"]}],
            "options": {"flexible_regions": ["anywhere_in_continent",
                                             "anywhere_in_world"]},
        }],
        "workplaceTypes": ["Remote"],
        "defaultToUserLocation": False, "userLocation": None,
        "currency": {"label": "Any", "value": None},
        "frequency": {"label": "Any", "value": None},
        "calcFrequency": "Yearly",
        "commitmentTypes": ["Full Time", "Part Time", "Contract", "Internship",
                            "Temporary", "Seasonal", "Volunteer"],
        "jobTitleQuery": "", "jobDescriptionQuery": "",
        "seniorityLevel": ["No Prior Experience Required", "Entry Level",
                           "Mid Level", "Senior Level"],
        "roleTypes": ["Individual Contributor", "People Manager"],
        "roleYoeRange": [0, 20], "managementYoeRange": [0, 20],
        "searchQuery": keywords or "",
        "dateFetchedPastNDays": 61, "sortBy": "default",
        "companyPublicOrPrivate": "all", "isNonProfit": "all",
        "companyNames": [], "excludedCompanyNames": [],
        "industries": [], "excludedIndustries": [],
        "departments": [], "restrictedSearchAttributes": [],
    }


def _extract_rows(data):
    rows = []
    if isinstance(data, dict):
        for k in ("results", "jobs", "data", "items", "content"):
            if isinstance(data.get(k), list):
                return data[k]
        if isinstance(data.get("hits"), dict):
            return [h.get("_source", h) for h in data["hits"].get("hits", [])]
    elif isinstance(data, list):
        return data
    return rows


def _hiringcafe_fetch(keywords):
    """Return (rows, info). Tries the full payload, then a minimal one.
    `info` is a small diagnostic used by /api/diag/hiringcafe."""
    info = {"attempts": []}
    states = [("full", _hiringcafe_full_state(keywords)),
              ("minimal", {"searchQuery": keywords or "",
                           "workplaceTypes": ["Remote"],
                           "dateFetchedPastNDays": 61, "sortBy": "default"})]
    for name, state in states:
        a = {"payload": name}
        try:
            r = requests.post(_HIRINGCAFE_URL, headers=_hiringcafe_headers(),
                              json={"size": 100, "page": 0, "searchState": state},
                              timeout=30)
            a["status"] = r.status_code
            if r.status_code == 200:
                data = r.json()
                rows = _extract_rows(data)
                a["rows"] = len(rows)
                a["top_keys"] = (list(data.keys())[:12]
                                 if isinstance(data, dict) else "list")
                if rows and isinstance(rows[0], dict):
                    a["first_row_keys"] = list(rows[0].keys())[:15]
                info["attempts"].append(a)
                if rows:
                    return rows, info
            else:
                a["body"] = r.text[:300]
                info["attempts"].append(a)
        except Exception as e:
            a["error"] = str(e)[:300]
            info["attempts"].append(a)
    return [], info


def _hiringcafe_build(row):
    if not isinstance(row, dict):
        return None
    v5 = row.get("v5_processed_job_data") or {}
    ji = row.get("job_information") or {}
    title = (v5.get("core_job_title") or v5.get("job_title")
             or ji.get("title") or row.get("title") or row.get("job_title") or "")
    if not title:
        return None
    company = (v5.get("company_name") or row.get("company_name")
               or ji.get("company_name") or row.get("company") or "")
    loc = v5.get("formatted_workplace_location") or row.get("location") or ""
    if not loc:
        wc = v5.get("workplace_countries") or []
        loc = ", ".join(wc) if isinstance(wc, list) else str(wc)
    wtype = str(v5.get("workplace_type") or row.get("workplace_type") or "")
    region = loc or "Remote"
    if "remote" in wtype.lower() and "remote" not in region.lower():
        region = (region + " (Remote)").strip()
    desc = _strip_html(ji.get("description") or v5.get("requirements_summary")
                       or v5.get("role_activities_summary") or row.get("description") or "")
    url = (row.get("apply_url") or ji.get("apply_url") or row.get("url")
           or ji.get("url") or "https://hiring.cafe/")
    smin = _to_int(v5.get("yearly_min_compensation") or v5.get("min_compensation")
                   or row.get("salary_min"))
    smax = _to_int(v5.get("yearly_max_compensation") or v5.get("max_compensation")
                   or row.get("salary_max"))
    sraw = f"${smin or '?'}-${smax or '?'}" if (smin or smax) else ""
    jb = _blank_job()
    jb.update(source="hiring.cafe", title=title, company=str(company),
              region=str(region), url=url, description=desc,
              posted=str(v5.get("estimated_publish_date", "")),
              salary_min=smin, salary_max=smax, salary_raw=sraw)
    return jb


def fetch_hiringcafe(keywords):
    """hiring.cafe — large ATS aggregator via its real search API."""
    out = []
    try:
        rows, info = _hiringcafe_fetch(keywords)
        for row in rows:
            jb = _hiringcafe_build(row)
            if jb and _matches(keywords, jb["title"], jb["description"]):
                out.append(jb)
        if not out and info["attempts"]:
            app.logger.warning("hiring.cafe: 0 jobs. diag=%s", info["attempts"])
    except Exception as e:
        app.logger.warning("hiring.cafe failed: %s", e)
    return out


def _find_job_list(obj, best=None):
    """Recursively find the largest list of job-like dicts inside a nested
    structure (used to parse Next.js __NEXT_DATA__ blobs)."""
    if best is None:
        best = []
    if isinstance(obj, list):
        joblike = [x for x in obj if isinstance(x, dict) and
                   any(k in x for k in ("title", "position", "jobTitle", "name")) and
                   any(k in x for k in ("company", "company_name", "companyName",
                                        "url", "slug", "link"))]
        if len(joblike) > len(best):
            best = joblike
        for x in obj:
            best = _find_job_list(x, best)
    elif isinstance(obj, dict):
        for v in obj.values():
            best = _find_job_list(v, best)
    return best


def fetch_remoteyeah(keywords):
    """RemoteYeah — Next.js remote board. Parses its embedded __NEXT_DATA__
    JSON (no public API). Best-effort; falls back to nothing on layout changes."""
    out = []
    urls = ["https://remoteyeah.com/", "https://remoteyeah.com/remote-jobs-in-latin-america"]
    for page_url in urls:
        try:
            r = requests.get(page_url, headers=HTTP_HEADERS, timeout=20)
            r.raise_for_status()
            m = re.search(r'<script id="__NEXT_DATA__"[^>]*>(.*?)</script>',
                          r.text, re.S)
            if not m:
                continue
            data = json.loads(m.group(1))
            for j in _find_job_list(data):
                title = (j.get("title") or j.get("position")
                         or j.get("jobTitle") or j.get("name") or "")
                if not title:
                    continue
                desc = _strip_html(j.get("description") or j.get("excerpt")
                                   or j.get("summary") or "")
                if not _matches(keywords, title, desc):
                    continue
                comp = j.get("company") or j.get("company_name") or j.get("companyName") or ""
                if isinstance(comp, dict):
                    comp = comp.get("name", "")
                region = (j.get("location") or j.get("region")
                          or j.get("candidate_location") or "")
                if "latin" in page_url and not region:
                    region = "Latin America"
                region = region or "Remote"
                slug = j.get("url") or j.get("slug") or j.get("link") or ""
                if slug and not str(slug).startswith("http"):
                    slug = "https://remoteyeah.com/" + str(slug).lstrip("/")
                smin = _to_int(j.get("salary_min") or j.get("minSalary"))
                smax = _to_int(j.get("salary_max") or j.get("maxSalary"))
                sraw = f"${smin or '?'}-${smax or '?'}" if (smin or smax) else ""
                jb = _blank_job()
                jb.update(source="RemoteYeah", title=title, company=str(comp),
                          region=str(region), url=slug, description=desc,
                          salary_min=smin, salary_max=smax, salary_raw=sraw)
                out.append(jb)
            if out:
                break
        except Exception as e:
            app.logger.warning("RemoteYeah failed (%s): %s", page_url, e)
    # dedupe within RemoteYeah by title+company
    seen, uniq = set(), []
    for j in out:
        k = (j["title"].lower(), j["company"].lower())
        if k not in seen:
            seen.add(k); uniq.append(j)
    return uniq


SOURCES = {
    "remotive": fetch_remotive,
    "remoteok": fetch_remoteok,
    "arbeitnow": fetch_arbeitnow,
    "jobicy": fetch_jobicy,
    "himalayas": fetch_himalayas,
    "themuse": fetch_themuse,
    "weworkremotely": fetch_weworkremotely,
    "workingnomads": fetch_workingnomads,
    "jobspresso": fetch_jobspresso,
    "getonbrd": fetch_getonbrd,
}
DEFAULT_SOURCES = ["remotive", "remoteok", "arbeitnow", "jobicy",
                   "himalayas", "workingnomads", "getonbrd"]


def aggregate_jobs(keywords, sources):
    """Fetch from selected sources concurrently, dedupe, assign ids."""
    results = []
    threads = []

    def run(fn):
        try:
            results.extend(fn(keywords))
        except Exception as e:
            app.logger.warning("source error: %s", e)

    for s in sources:
        fn = SOURCES.get(s)
        if fn:
            t = threading.Thread(target=run, args=(fn,))
            t.start()
            threads.append(t)
    for t in threads:
        t.join(timeout=25)

    seen = set()
    deduped = []
    for j in results:
        key = (j["title"].strip().lower(), j["company"].strip().lower())
        if not j["title"].strip() or key in seen:
            continue
        seen.add(key)
        j["id"] = abs(hash(key)) % (10 ** 9)
        deduped.append(j)
    return deduped


def query_relevance(job, keywords):
    """Score how well a job matches the typed keywords (0 = no match).
    Title hits weigh most, then tags/category, then description."""
    kw = _kwlist(keywords)
    if not kw:
        return 1.0
    title = job["title"].lower()
    tagcat = (job["category"] + " " + " ".join(job["tags"])).lower()
    desc = job["description"][:1500].lower()
    phrase = " ".join(kw)
    score = 0.0
    if len(kw) > 1 and phrase in title:
        score += 6
    for t in kw:
        if t in title:
            score += 3
        elif t in tagcat:
            score += 2
        elif t in desc:
            score += 1
    hits = sum(1 for t in kw if (t in title or t in tagcat or t in desc))
    score += 1.5 * (hits / len(kw))   # reward covering more of the query
    return score


def query_is_relevant(job, keywords):
    """Keep a job only if the query actually appears in it (drops the loose,
    category-only matches some sources return)."""
    kw = _kwlist(keywords)
    if not kw:
        return True
    title = job["title"].lower()
    tagcat = (job["category"] + " " + " ".join(job["tags"])).lower()
    desc = job["description"][:1500].lower()
    return any((t in title or t in tagcat or t in desc) for t in kw)


# ---------------------------------------------------------------------------
# Resume parsing, work-experience extraction + skills
# ---------------------------------------------------------------------------
STOPWORDS = set("""a an the and or but for nor so yet of to in on at by with from as is are
was were be been being this that these those it its i you he she they we my our your their
will would can could should may might must have has had do does did not no your about into
over under again further then once here there all any both each few more most other some such""".split())

SKILL_LEXICON = [
    "python", "javascript", "typescript", "java", "c++", "c#", "go", "golang",
    "ruby", "php", "rust", "scala", "kotlin", "swift", "sql", "nosql", "react",
    "angular", "vue", "node", "node.js", "django", "flask", "fastapi", "spring",
    "rails", ".net", "express", "next.js", "aws", "azure", "gcp", "docker",
    "kubernetes", "terraform", "ci/cd", "jenkins", "git", "linux", "postgres",
    "postgresql", "mysql", "mongodb", "redis", "kafka", "spark", "hadoop",
    "tableau", "power bi", "excel", "pandas", "numpy", "pytorch", "tensorflow",
    "machine learning", "deep learning", "nlp", "data science", "data analysis",
    "etl", "airflow", "snowflake", "dbt", "salesforce", "sap", "hubspot",
    "marketing", "seo", "sem", "content", "copywriting", "social media",
    "project management", "scrum", "agile", "kanban", "jira", "product management",
    "ux", "ui", "figma", "sketch", "adobe", "photoshop", "illustrator",
    "customer success", "customer support", "sales", "account management",
    "business development", "finance", "accounting", "bookkeeping", "payroll",
    "recruiting", "hr", "operations", "logistics", "qa", "testing", "selenium",
    "cypress", "graphql", "rest", "api", "microservices", "devops", "sre",
    "security", "cybersecurity", "blockchain", "solidity", "english", "spanish",
    "portuguese", "bilingual", "leadership", "communication", "ansible",
    "prometheus", "grafana", "helm", "argocd", "gitlab", "github actions",
    "bash", "powershell", "vmware", "networking", "cloud",
]

_SKILL_PATTERNS = [
    (sk, re.compile(r"(?<![a-z0-9+#])" + re.escape(sk) + r"(?![a-z0-9+#])"))
    for sk in SKILL_LEXICON
]


def extract_skills(text):
    low = (text or "").lower()
    return sorted({sk for sk, pat in _SKILL_PATTERNS if pat.search(low)})


def tokenize(text):
    toks = re.findall(r"[a-zA-Z][a-zA-Z+.#]{1,}", (text or "").lower())
    return [t for t in toks if t not in STOPWORDS and len(t) > 2]


# Section headers used to isolate work experience from the rest of the resume.
EXP_HEADERS = [
    "work experience", "professional experience", "experience",
    "employment history", "employment", "work history", "career history",
    "relevant experience", "professional background", "career summary",
    "experiencia laboral", "experiencia profesional", "experiencia",
]
# Sections we EXCLUDE from matching (courses, schooling, etc.).
EXCLUDE_HEADERS = [
    "education", "academic", "academics", "courses", "course work",
    "coursework", "certifications", "certification", "certificates",
    "certificate", "licenses", "training", "trainings", "awards", "honors",
    "publications", "references", "interests", "hobbies", "volunteer",
    "volunteering", "extracurricular", "activities", "educacion", "educación",
    "cursos", "certificaciones", "formacion", "formación",
]
# Sections that ALSO count as real competencies (kept for matching).
KEEP_HEADERS = [
    "skills", "technical skills", "core competencies", "competencies",
    "summary", "professional summary", "profile", "about", "languages",
    "language", "tools", "technologies", "tech stack", "habilidades",
    "resumen", "perfil",
]


def _header_kind(line):
    s = line.strip().strip(":").strip().lower()
    if not s or len(s) > 42 or len(s.split()) > 6:
        return None
    for h in EXCLUDE_HEADERS:
        if s == h or s.startswith(h):
            return "exclude"
    for h in EXP_HEADERS:
        if s == h or s.startswith(h):
            return "include"
    for h in KEEP_HEADERS:
        if s == h or s.startswith(h):
            return "include"
    return None


def extract_experience_text(text):
    """Return only the work-experience (+skills/summary) parts of a resume,
    dropping education / courses / certifications. Falls back to full text if
    the resume has no recognizable section headers."""
    lines = (text or "").splitlines()
    mode = "include"          # text before the first header (name/summary) is kept
    saw_header = False
    kept = []
    for line in lines:
        kind = _header_kind(line)
        if kind == "exclude":
            mode, saw_header = "exclude", True
            continue
        if kind == "include":
            mode, saw_header = "include", True
            continue
        if mode == "include":
            kept.append(line)
    result = "\n".join(kept).strip()
    # Fall back to the full resume only if there were no recognizable section
    # headers at all, or we captured almost nothing.
    if not saw_header or len(result) < 40:
        return text or ""
    return result


def parse_resume_bytes(data, filename):
    name = (filename or "").lower().strip()
    if name.endswith(".pdf"):
        if not pdfplumber:
            raise RuntimeError("PDF support not installed. Run: pip install pdfplumber")
        text = []
        try:
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for page in pdf.pages:
                    text.append(page.extract_text() or "")
        except Exception:
            raise RuntimeError("Could not read this PDF — it may be corrupted or password-protected.")
        t = "\n".join(text).strip()
        if not t:
            raise RuntimeError(
                "This PDF has no selectable text (likely a scanned image). "
                "Export a text-based PDF, or upload your resume as .docx.")
        return t
    if name.endswith(".docx"):
        if not docx:
            raise RuntimeError("DOCX support not installed. Run: pip install python-docx")
        try:
            d = docx.Document(io.BytesIO(data))
        except Exception:
            raise RuntimeError(
                "Could not open this .docx (it may be corrupted, or actually a .doc "
                "renamed). Open it in Word and use Save As -> .docx or PDF.")
        parts = [p.text for p in d.paragraphs]
        for tbl in d.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts).strip()
    if name.endswith(".txt"):
        return data.decode("utf-8", errors="ignore").strip()
    if name.endswith(".doc"):
        raise RuntimeError(
            "Old .doc format isn't supported. Open it in Word and Save As -> .docx or PDF.")
    raise RuntimeError("Unsupported file type — upload a .pdf, .docx or .txt resume.")


# ---------------------------------------------------------------------------
# Matching (local) — uses ONLY the work-experience text
# ---------------------------------------------------------------------------
def local_match_score(experience_text, resume_skills, job):
    job_blob = f"{job['title']} {job['category']} {' '.join(job['tags'])} {job['description'][:1500]}".lower()
    job_skills = set(extract_skills(job_blob))
    res_skills = set(resume_skills)

    matched = sorted(res_skills & job_skills)
    gaps = sorted(job_skills - res_skills)

    skill_component = (len(matched) / len(job_skills)) if job_skills else 0.0

    res_tokens = set(tokenize(experience_text))
    job_tokens = tokenize(job_blob)
    overlap = (sum(1 for t in job_tokens if t in res_tokens) / len(job_tokens)) if job_tokens else 0.0

    title_low = job["title"].lower()
    title_boost = 0.15 if any(s in title_low for s in res_skills) else 0.0

    score = 100 * (0.55 * skill_component + 0.30 * overlap + title_boost)
    return max(0, min(100, round(score))), matched, gaps[:12]


# ---------------------------------------------------------------------------
# LLM client (optional)
# ---------------------------------------------------------------------------
def llm_available():
    s = STATE["settings"]
    if s["llm_provider"] == "free":
        return True   # keyless hosted engine (Pollinations) — no login required
    return s["llm_provider"] in ("openai", "anthropic") and bool(s["api_key"])


def _llm_free(system, user, max_tokens, want_json):
    """Free, keyless LLM via Pollinations (OpenAI-compatible). No API key/login.
    Note: this sends the prompt to a shared free public service."""
    import urllib.parse
    try:
        r = requests.post(
            "https://text.pollinations.ai/openai",
            headers={"Content-Type": "application/json"},
            json={"model": "openai",
                  "messages": [{"role": "system", "content": system},
                               {"role": "user", "content": user}],
                  "temperature": 0.4, "max_tokens": max_tokens, "private": True},
            timeout=90)
        r.raise_for_status()
        data = r.json()
        return data["choices"][0]["message"]["content"]
    except Exception:
        # Fallback to the plain GET prompt endpoint.
        prompt = (system + "\n\n" + user)[:1800]
        r = requests.get("https://text.pollinations.ai/" + urllib.parse.quote(prompt),
                         params={"model": "openai"}, timeout=90)
        r.raise_for_status()
        return r.text


def llm_chat(system, user, max_tokens=900, want_json=False):
    s = STATE["settings"]
    provider, key = s["llm_provider"], s["api_key"]
    if provider == "free":
        return _llm_free(system, user, max_tokens, want_json)
    if provider == "openai":
        body = {
            "model": s["openai_model"],
            "messages": [{"role": "system", "content": system},
                         {"role": "user", "content": user}],
            "temperature": 0.4, "max_tokens": max_tokens,
        }
        if want_json:
            body["response_format"] = {"type": "json_object"}
        r = requests.post("https://api.openai.com/v1/chat/completions",
                          headers={"Authorization": f"Bearer {key}",
                                   "Content-Type": "application/json"},
                          json=body, timeout=60)
        r.raise_for_status()
        return r.json()["choices"][0]["message"]["content"]
    elif provider == "anthropic":
        r = requests.post("https://api.anthropic.com/v1/messages",
                          headers={"x-api-key": key,
                                   "anthropic-version": "2023-06-01",
                                   "Content-Type": "application/json"},
                          json={"model": s["anthropic_model"], "max_tokens": max_tokens,
                                "system": system,
                                "messages": [{"role": "user", "content": user}]},
                          timeout=60)
        r.raise_for_status()
        return r.json()["content"][0]["text"]
    raise RuntimeError("No LLM provider configured")


# ---------------------------------------------------------------------------
# Suggestions
# ---------------------------------------------------------------------------
def resume_is_strong(score, gaps):
    """Heuristic: resume already fits the role, so suggest no changes."""
    return score >= 80 and len(gaps) == 0


def local_tips(experience_text, job, matched, gaps):
    tips = []
    if gaps:
        tips.append("Consider adding these role-relevant skills if you have real "
                    "experience with them: " + ", ".join(gaps[:8]) + ".")
    tips.append("Mirror the exact job title \"%s\" near the top so ATS scans rank "
                "you higher." % job["title"])
    if matched:
        tips.append("Lead your most relevant work-experience bullets with: "
                    + ", ".join(matched[:6]) + ".")
    tips.append("Quantify 2-3 achievements with metrics (%, $, time saved) tied to "
                "this role's responsibilities.")
    if any(w in job["description"].lower() for w in ["english", "fluent", "bilingual"]):
        tips.append("Make sure English proficiency is visible near the top.")
    return tips


def llm_suggestions(full_resume, job):
    system = ("You are an expert technical recruiter and resume writer who helps "
              "LatAm-based candidates win remote roles. Base everything on the "
              "candidate's WORK EXPERIENCE, not courses/education. If the resume "
              "already fits the role well, say so and propose no edits. Never "
              "invent experience the candidate does not have.")
    user = ("JOB TITLE: " + job["title"] + "\nCOMPANY: " + job["company"] +
            "\nJOB DESCRIPTION (truncated):\n" + job["description"][:2500] +
            "\n\nCANDIDATE RESUME (verbatim, truncated):\n" + full_resume[:6000] +
            "\n\nReturn ONLY JSON with keys:\n"
            "'no_changes_needed' (boolean: true if the resume is already a strong fit),\n"
            "'assessment' (1-2 sentence verdict),\n"
            "'edits' (array of {\"find\",\"replace\",\"reason\"} where 'find' is an EXACT "
            "verbatim substring copied from the resume to replace, and 'replace' is the "
            "improved wording; keep them short and sentence-level, only where they add value),\n"
            "'keywords_to_add' (array of strings present in the job but missing from the resume),\n"
            "'rewritten_summary' (optional 2-3 sentence tailored summary),\n"
            "'tips' (array of short manual suggestions you could not express as exact edits).")
    raw = llm_chat(system, user, max_tokens=1100, want_json=True)
    try:
        return json.loads(raw)
    except Exception:
        m = re.search(r"\{.*\}", raw, re.S)
        return json.loads(m.group(0)) if m else {"tips": [raw]}


# ---------------------------------------------------------------------------
# Tailored docx export
# ---------------------------------------------------------------------------
def build_tailored_docx(resume_text, job, summary, suggestions, keywords):
    if not docx:
        raise RuntimeError("python-docx not installed")
    from docx.shared import Pt, RGBColor
    d = docx.Document()
    d.add_heading("Tailored Resume", level=0)
    sub = d.add_paragraph()
    run = sub.add_run(f"Optimized for: {job['title']} - {job['company']}")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    d.add_heading("Professional Summary", level=1)
    d.add_paragraph(summary or "")

    if keywords:
        d.add_heading("Key Skills (ATS-aligned)", level=1)
        d.add_paragraph(" - ".join(keywords))

    d.add_heading("Suggested Improvements Applied", level=1)
    for s in suggestions:
        d.add_paragraph(style="List Bullet").add_run(s)

    d.add_heading("Original Resume Content", level=1)
    for line in resume_text.splitlines():
        if line.strip():
            d.add_paragraph(line.strip())

    fr = d.add_paragraph().add_run(
        f"\nGenerated by JobHunt LatAm on {datetime.now():%Y-%m-%d %H:%M}")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    out = io.BytesIO()
    d.save(out)
    out.seek(0)
    return out


# ---------------------------------------------------------------------------
# In-place editing of the ORIGINAL .docx (preserves all formatting)
# ---------------------------------------------------------------------------
def _iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _replace_in_paragraph(par, find, repl):
    """Replace `find` with `repl` inside a paragraph, keeping its formatting."""
    if not find or find not in par.text:
        return False
    # Easy case: the text lives inside a single run -> formatting is untouched.
    for run in par.runs:
        if find in run.text:
            run.text = run.text.replace(find, repl)
            return True
    # Spans several runs: rebuild onto the first run (keeps paragraph style:
    # bullet, indent, alignment, and the first run's character formatting).
    runs = par.runs
    if not runs:
        return False
    full = "".join(r.text for r in runs)
    if find not in full:
        return False
    runs[0].text = full.replace(find, repl)
    for r in runs[1:]:
        r.text = ""
    return True


def apply_edits_to_docx(original_bytes, edits, summary, insert_summary, keywords):
    """Open the user's original .docx and apply surgical edits in place."""
    d = docx.Document(io.BytesIO(original_bytes))
    applied = 0
    for e in edits or []:
        f = (e.get("find") or "").strip()
        r = e.get("replace") or ""
        if not f:
            continue
        for p in _iter_paragraphs(d):
            if _replace_in_paragraph(p, f, r):
                applied += 1
                break
    # Append missing keywords to an existing Skills/Tools line (if any).
    if keywords:
        for p in _iter_paragraphs(d):
            low = p.text.lower()
            if any(h in low for h in ("skills", "technologies", "tech stack", "tools")):
                target = p.runs[-1] if p.runs else p.add_run("")
                tail = target.text.rstrip()
                sep = "" if tail.endswith((",", ":", ";")) else ", "
                target.text = target.text + sep + ", ".join(keywords)
                break
    # Optionally drop a tailored summary line at the very top (non-destructive).
    if summary and insert_summary and d.paragraphs:
        d.paragraphs[0].insert_paragraph_before(summary)
    out = io.BytesIO()
    d.save(out)
    out.seek(0)
    return out, applied


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------
@app.route("/")
def index():
    return render_template("index.html")


@app.route("/favicon.ico")
def favicon():
    return Response(status=204)


@app.route("/api/countries")
def countries():
    return jsonify({"countries": [c.title() for c in LATAM_COUNTRIES]})


@app.route("/api/diag/hiringcafe")
def diag_hiringcafe():
    """Visit /api/diag/hiringcafe?q=devops to see exactly what hiring.cafe
    returns on YOUR machine (status, keys, counts) — helpful for debugging."""
    q = request.args.get("q", "devops")
    rows, info = _hiringcafe_fetch(q)
    parsed = [b for b in (_hiringcafe_build(r) for r in rows[:5]) if b]
    return jsonify({
        "query": q,
        "diagnostics": info,
        "rows_found": len(rows),
        "sample_parsed": [{"title": p["title"], "company": p["company"],
                           "region": p["region"]} for p in parsed],
    })


@app.route("/api/settings", methods=["GET", "POST"])
def settings():
    if request.method == "POST":
        body = request.get_json(force=True)
        with _LOCK:
            s = STATE["settings"]
            s["llm_provider"] = body.get("llm_provider", s["llm_provider"])
            if "api_key" in body:
                s["api_key"] = body["api_key"]
            s["openai_model"] = body.get("openai_model", s["openai_model"])
            s["anthropic_model"] = body.get("anthropic_model", s["anthropic_model"])
        return jsonify({"ok": True, "llm_available": llm_available()})
    s = STATE["settings"]
    return jsonify({"llm_provider": s["llm_provider"], "has_key": bool(s["api_key"]),
                    "openai_model": s["openai_model"], "anthropic_model": s["anthropic_model"],
                    "llm_available": llm_available()})


@app.route("/api/search", methods=["POST"])
def search():
    body = request.get_json(force=True)
    keywords = (body.get("keywords") or "").strip()
    sources = body.get("sources") or DEFAULT_SOURCES
    min_salary = body.get("min_salary")
    latam_only = bool(body.get("latam_only", True))
    sort_by = body.get("sort_by", "match")
    country = _norm_country(body.get("country", "all"))

    jobs = aggregate_jobs(keywords, sources)

    # Keyword relevance: drop loose/unrelated matches and score the rest so the
    # most on-topic roles rank first (esp. important when no resume is loaded).
    if keywords:
        jobs = [j for j in jobs if query_is_relevant(j, keywords)]
    for j in jobs:
        j["_rel"] = query_relevance(j, keywords)

    # Eligibility per the selected country.
    for j in jobs:
        blob = f"{j['region']} {j['description'][:600]} {' '.join(j['tags'])}"
        label, score, reason = latam_eligibility(blob, j["region"], country)
        j["latam_label"], j["latam_score"], j["latam_reason"] = label, score, reason

    if min_salary:
        try:
            thr = int(min_salary)
            kept = []
            for j in jobs:
                if j["salary_max"] is None and j["salary_min"] is None:
                    kept.append(j)
                else:
                    top = j["salary_max"] or j["salary_min"]
                    if top and top >= thr:
                        kept.append(j)
            jobs = kept
        except (TypeError, ValueError):
            pass

    if latam_only:
        jobs = [j for j in jobs if j["latam_score"] >= 50]

    have_resume = bool(STATE["resume_experience"] or STATE["resume_text"])
    exp_text = STATE["resume_experience"] or STATE["resume_text"]
    for j in jobs:
        if have_resume:
            score, matched, gaps = local_match_score(exp_text, STATE["resume_skills"], j)
            j["match_score"], j["matched_skills"], j["skill_gaps"] = score, matched, gaps
        else:
            j["match_score"], j["matched_skills"], j["skill_gaps"] = None, [], []

    if sort_by == "salary":
        jobs.sort(key=lambda x: (x["salary_max"] or x["salary_min"] or 0), reverse=True)
    elif sort_by == "latam":
        jobs.sort(key=lambda x: (x["latam_score"], x.get("_rel", 0)), reverse=True)
    else:  # "match" (default): resume fit first, then keyword relevance
        if have_resume:
            jobs.sort(key=lambda x: (x["match_score"] or 0, x.get("_rel", 0),
                                     x["latam_score"]), reverse=True)
        else:
            jobs.sort(key=lambda x: (x.get("_rel", 0), x["latam_score"]), reverse=True)

    with _LOCK:
        STATE["last_jobs"] = jobs

    return jsonify({"count": len(jobs), "have_resume": have_resume, "jobs": jobs[:150]})


@app.route("/api/upload_resume", methods=["POST"])
def upload_resume():
    if "file" not in request.files:
        return jsonify({"error": "No file received."}), 400
    f = request.files["file"]
    if not f or not f.filename:
        return jsonify({"error": "Empty filename."}), 400
    try:
        data = f.read()
        text = parse_resume_bytes(data, f.filename)
    except Exception as e:
        app.logger.warning("resume parse failed (%s): %s", f.filename, e)
        return jsonify({"error": str(e)}), 400
    if not text or not text.strip():
        return jsonify({"error": "Could not extract any text from that file."}), 400
    experience = extract_experience_text(text)
    skills = extract_skills(experience)
    is_docx = (f.filename or "").lower().endswith(".docx")
    with _LOCK:
        STATE["resume_text"] = text
        STATE["resume_experience"] = experience
        STATE["resume_filename"] = f.filename
        STATE["resume_skills"] = skills
        STATE["resume_docx"] = data if is_docx else None
    return jsonify({"ok": True, "filename": f.filename, "chars": len(text),
                    "experience_chars": len(experience), "skills": skills,
                    "can_edit_format": is_docx, "preview": experience[:600]})


@app.route("/api/suggest", methods=["POST"])
def suggest():
    body = request.get_json(force=True)
    job_id = body.get("job_id")
    if not (STATE["resume_experience"] or STATE["resume_text"]):
        return jsonify({"error": "Upload a resume first."}), 400
    job = next((j for j in STATE["last_jobs"] if j["id"] == job_id), None)
    if not job:
        return jsonify({"error": "Job not found. Re-run the search."}), 404

    exp_text = STATE["resume_experience"] or STATE["resume_text"]
    score, matched, gaps = local_match_score(exp_text, STATE["resume_skills"], job)
    strong = resume_is_strong(score, gaps)
    result = {
        "job_id": job_id, "title": job["title"], "company": job["company"],
        "match_score": score, "matched_skills": matched, "skill_gaps": gaps,
        "engine": "local",
        "no_changes_needed": strong,
        "assessment": ("Your resume already covers this role's core requirements "
                       "— no changes recommended."
                       if strong else
                       "A few targeted tweaks could strengthen your fit."),
        "can_edit_format": bool(STATE["resume_docx"]),
        "summary": "", "edits": [],
        "keywords_to_add": ([] if strong else gaps[:10]),
        "tips": ([] if strong else local_tips(exp_text, job, matched, gaps)),
    }
    if llm_available():
        try:
            data = llm_suggestions(STATE["resume_text"], job)
            result["engine"] = STATE["settings"]["llm_provider"]
            if "no_changes_needed" in data:
                result["no_changes_needed"] = bool(data["no_changes_needed"])
            if data.get("assessment"):
                result["assessment"] = data["assessment"]
            result["edits"] = data.get("edits", []) or []
            result["tips"] = data.get("tips", []) or []
            result["summary"] = data.get("rewritten_summary", "") or ""
            if data.get("keywords_to_add"):
                result["keywords_to_add"] = data["keywords_to_add"]
            if result["no_changes_needed"]:
                result["edits"], result["tips"], result["keywords_to_add"] = [], [], []
        except Exception as e:
            result["llm_error"] = str(e)
    return jsonify(result)


@app.route("/api/apply", methods=["POST"])
def apply_changes():
    body = request.get_json(force=True)
    job_id = body.get("job_id")
    summary = body.get("summary", "")
    insert_summary = bool(body.get("insert_summary", False))
    edits = body.get("edits", []) or []
    keywords = body.get("keywords_to_add", []) or []
    tips = body.get("suggestions", []) or []
    if not STATE["resume_text"]:
        return jsonify({"error": "Upload a resume first."}), 400
    job = next((j for j in STATE["last_jobs"] if j["id"] == job_id), None)
    if not job:
        return jsonify({"error": "Job not found."}), 404
    safe = re.sub(r"[^a-zA-Z0-9]+", "_", f"{job['company']}_{job['title']}")[:50]
    fname = f"Tailored_Resume_{safe}.docx"
    try:
        if STATE["resume_docx"]:
            # Edit the ORIGINAL .docx -> keeps the exact layout/format intact.
            buf, _applied = apply_edits_to_docx(
                STATE["resume_docx"], edits, summary, insert_summary, keywords)
        else:
            # PDF/TXT upload: no original Word formatting to preserve.
            if not summary:
                sk = ", ".join(STATE["resume_skills"][:6]) or "relevant experience"
                summary = (f"Results-driven professional targeting the {job['title']} "
                           f"role at {job['company']}, with strengths in {sk}.")
            buf = build_tailored_docx(STATE["resume_text"], job, summary, tips, keywords)
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    return send_file(buf,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     as_attachment=True, download_name=fname)


@app.route("/api/status")
def status():
    return jsonify({"resume_loaded": bool(STATE["resume_text"]),
                    "resume_filename": STATE["resume_filename"],
                    "resume_skills": STATE["resume_skills"],
                    "llm_available": llm_available(),
                    "last_job_count": len(STATE["last_jobs"])})


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    print("\n  JobHunt LatAm running at  http://127.0.0.1:%d\n" % port)
    app.run(host="127.0.0.1", port=port, debug=False)
